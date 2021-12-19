import datetime
import os

import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from scipy import signal
from scipy.integrate import simps

from topograph import get_psds_alpha, get_psds_beta, get_psds_theta, plot_topomap

matplotlib.use("Agg")


def neuro(input_data, input_name):
    input_data = os.getcwd() + "/out/" + input_data
    print("input_data: " + input_data)
    print("input_name: " + input_name)
    df = pd.read_table(
        input_data,
        header=None,
        delim_whitespace=True,
        names=(
            "Date1",
            "Date2",
            "Fp1",
            "Fp2",
            "F7",
            "F8",
            "C3",
            "C4",
            "T3",
            "T4",
            "T5",
            "T6",
            "O1",
            "O2",
            "X1",
            "Other1",
            "Other2",
        ),
    )

    # dfの列を並び替える
    df = df[
        [
            "Date1",
            "Date2",
            "Fp1",
            "Fp2",
            "C3",
            "C4",
            "O1",
            "O2",
            "T3",
            "T4",
            "X1",
            "F7",
            "F8",
            "T5",
            "T6",
            "Other1",
            "Other2",
        ]
    ]

    df.drop(columns=["Date1", "Date2", "X1", "Other1", "Other2"], inplace=True)
    # 欠損値が１つでもある行を削除する
    df = df.dropna(how="any")

    # topomapを作成して保存する
    df_copy = df.T
    df_np = df_copy.values

    # シータ波のスペクトルを求める
    pwrs_theta, _ = get_psds_theta(df_np)

    # アルファ波のスペクトルを求める
    pwrs_alpha, _ = get_psds_alpha(df_np)

    # ベータ波のスペクトルを求める
    pwrs_beta, _ = get_psds_beta(df_np)

    # Zスコアを求める
    eeg_list = [
        "Fp1",
        "Fp2",
        "C3",
        "C4",
        "O1",
        "O2",
        "T3",
        "T4",
        "F7",
        "F8",
        "T5",
        "T6",
    ]

    npy_dir = os.getcwd() + "/npy/"
    np_load_dataset = [
        [
            npy_dir + "pwrs_rel_alpha_fp1_eval+train.npy",
            npy_dir + "pwrs_rel_beta_fp1_eval+train.npy",
            npy_dir + "pwrs_rel_theta_fp1_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_fp2_eval+train.npy",
            npy_dir + "pwrs_rel_beta_fp2_eval+train.npy",
            npy_dir + "pwrs_rel_theta_fp2_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_c3_eval+train.npy",
            npy_dir + "pwrs_rel_beta_c3_eval+train.npy",
            npy_dir + "pwrs_rel_theta_c3_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_c4_eval+train.npy",
            npy_dir + "pwrs_rel_beta_c4_eval+train.npy",
            npy_dir + "pwrs_rel_theta_c4_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_O1_eval+train.npy",
            npy_dir + "pwrs_rel_beta_O1_eval+train.npy",
            npy_dir + "pwrs_rel_theta_O1_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_O2_eval+train.npy",
            npy_dir + "pwrs_rel_beta_O2_eval+train.npy",
            npy_dir + "pwrs_rel_theta_O2_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_t3_eval+train.npy",
            npy_dir + "pwrs_rel_beta_t3_eval+train.npy",
            npy_dir + "pwrs_rel_theta_t3_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_t4_eval+train.npy",
            npy_dir + "pwrs_rel_beta_t4_eval+train.npy",
            npy_dir + "pwrs_rel_theta_t4_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_f7_eval+train.npy",
            npy_dir + "pwrs_rel_beta_f7_eval+train.npy",
            npy_dir + "pwrs_rel_theta_f7_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_f8_eval+train.npy",
            npy_dir + "pwrs_rel_beta_f8_eval+train.npy",
            npy_dir + "pwrs_rel_theta_f8_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_t5_eval+train.npy",
            npy_dir + "pwrs_rel_beta_t5_eval+train.npy",
            npy_dir + "pwrs_rel_theta_t5_eval+train.npy",
        ],
        [
            npy_dir + "pwrs_rel_alpha_t6_eval+train.npy",
            npy_dir + "pwrs_rel_beta_t6_eval+train.npy",
            npy_dir + "pwrs_rel_theta_t6_eval+train.npy",
        ],
    ]

    j_list = ["アルファ", "ベータ", "シータ"]

    # result_listは結果を文章で出力するために用意したリスト
    result_list = []

    # ３行１６列のリストを作る。そしてアルファ波のtopomapならlist[0][l]をlでイテレートして表示する。numpy配列にしないといけないはず。
    # →そのやり方が分からなかったので３種類のリストを作る！
    result_alpha = []
    result_beta = []
    result_theta = []

    # ある電極での相対スペクトル密度のZ値を求める
    for i in range(len(eeg_list)):
        sf = 500.0
        df_analyze = df.iloc[:, i]
        df_analyze_np = df_analyze.values
        # time = np.arange(df_analyze_np.size) / sf

        # Welch's periodogramを求める
        win = 4 * sf
        freqs, psd = signal.welch(df_analyze_np, sf, nperseg=win)
        freq_res = freqs[1] - freqs[0]

        # ベータ波の相対スペクトル密度を求める
        idx_beta = np.logical_and(freqs >= 12, freqs <= 30)
        beta_power = simps(psd[idx_beta], dx=freq_res)

        # アルファ波の相対スペクトル密度を求める
        idx_alpha = np.logical_and(freqs >= 8, freqs <= 12)
        alpha_power = simps(psd[idx_alpha], dx=freq_res)

        # シータ波の相対スペクトル密度を求める
        idx_theta = np.logical_and(freqs >= 4, freqs <= 8)
        theta_power = simps(psd[idx_theta], dx=freq_res)

        total_power = theta_power + alpha_power + beta_power
        relative_alpha = alpha_power / total_power
        relative_beta = beta_power / total_power
        relative_theta = theta_power / total_power
        relative_list = [relative_alpha, relative_beta, relative_theta]

        # それぞれの電極で計算する
        for j in range(3):
            # 標準偏差を求める
            sample_spectrum = np.load(np_load_dataset[i][j])
            z = (relative_list[j] - np.mean(sample_spectrum)) / np.std(sample_spectrum)
            result_list.append(
                "{}電極の{}波の相対パワースペクトルのZ値は{}です。".format(
                    eeg_list[i], j_list[j], round(z, 2)
                )
            )

        # sample_spectrum_alpha = np.load(np_load_dataset[i][0])
        z1 = (relative_list[0] - np.mean(sample_spectrum)) / np.std(sample_spectrum)
        result_alpha.append(z1)

        # sample_spectrum_beta = np.load(np_load_dataset[i][1])
        z2 = (relative_list[1] - np.mean(sample_spectrum)) / np.std(sample_spectrum)
        result_beta.append(z2)

        # sample_spectrum_theta = np.load(np_load_dataset[i][2])
        z3 = (relative_list[2] - np.mean(sample_spectrum)) / np.std(sample_spectrum)
        result_theta.append(z3)

    out_dir = os.getcwd() + "/out/"

    # シータ波のtopomapを出力する
    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_theta, ax, fig)
    plt.title("theta_topomap")
    plt.savefig(out_dir + "theta_save_topomap.png")

    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_alpha, ax, fig)
    plt.title("alpha_topomap")
    plt.savefig(out_dir + "alpha_save_topomap.png")

    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_beta, ax, fig)
    plt.title("beta_topomap")
    plt.savefig(out_dir + "beta_save_topomap.png")

    dt_now = datetime.datetime.now()
    document = Document()
    document.add_heading("結果報告書")
    document.add_paragraph(" ")
    # document.add_paragraph(input_data[11:])
    # document.add_paragraph("ファイル名：")
    document.add_paragraph("作成日：" + dt_now.strftime("%Y年%m月%d日"))

    file_name = os.path.splitext(os.path.basename(input_data))[0] + ".m00"
    # print("file_name: ", file_name)
    # document.add_paragraph(file_name)
    document.add_paragraph("ファイル名：" + file_name)

    document.add_paragraph("インプット名：" + input_name)
    document.add_paragraph(" ")

    # document.add_paragraph(input_name)
    # document.add_paragraph("作成日：" + dt_now.strftime("%Y年%m月%d日"))
    # document.add_paragraph(dt_now.strftime("%Y年%m月%d日"))
    document.add_paragraph("あなたのアルファ波の分布は以下のようになります。")
    document.add_picture(out_dir + "alpha_save_topomap.png", width=Inches(3.5))
    document.add_paragraph("あなたのベータ波の分布は以下のようになります。")
    document.add_picture(out_dir + "beta_save_topomap.png", width=Inches(3.5))
    document.add_paragraph("あなたのシータ波の分布は以下のようになります。")
    document.add_picture(out_dir + "theta_save_topomap.png", width=Inches(3.5))

    document.add_paragraph(
        "あなたの脳波検査の結果からZスコアを算出します。（Zスコアは1297人の正常被験者のデータセットから算出しております）"
    )

    for k in range(len(result_list)):
        document.add_paragraph(" ")
        document.paragraphs[13 + k].add_run(result_list[k])

    print("282 input_data: ", input_data)
    document.save(
        os.path.splitext(
            os.path.dirname(input_data) + "/" + os.path.basename(input_data)
        )[0]
        + ".docx"
    )
    print("286 document saved to input_data")
